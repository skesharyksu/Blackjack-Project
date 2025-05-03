from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_word_report():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Blackjack AI: A Q-Learning Approach', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_heading('Final Project Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Introduction section
    doc.add_heading('1. Introduction', level=1)
    
    # Project Overview
    doc.add_heading('Project Overview', level=2)
    doc.add_paragraph('This project implements an AI agent that learns to play Blackjack using Q-learning, a model-free reinforcement learning algorithm. The system simulates the game environment, allowing the agent to learn optimal strategies through trial and error. The implementation includes a custom Blackjack environment, Q-learning algorithm, and tools for tracking and visualizing the learning process.')
    
    # Problem Statement
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph('Blackjack presents an interesting challenge for AI systems due to its combination of chance and strategy. While basic strategy exists, developing an AI that can learn and adapt its strategy through experience is valuable for both educational and practical purposes. The challenge lies in creating an agent that can learn optimal decisions without explicit programming of game rules or strategies.')
    
    # Scope and Challenges
    doc.add_heading('Scope and Challenges', level=2)
    scope = doc.add_paragraph('The project scope includes:')
    scope.add_run('\n• Implementing a custom Blackjack environment')
    scope.add_run('\n• Developing a Q-learning agent')
    scope.add_run('\n• Training the agent through simulation')
    scope.add_run('\n• Evaluating performance against basic strategy')
    
    challenges = doc.add_paragraph('Key challenges:')
    challenges.add_run('\n• State space complexity in Blackjack')
    challenges.add_run('\n• Balancing exploration and exploitation')
    challenges.add_run('\n• Efficient learning with sparse rewards')
    challenges.add_run('\n• Measuring and comparing performance')
    
    # Add Background section
    doc.add_heading('2. Background & Related Work', level=1)
    
    # Existing Approaches
    doc.add_heading('Existing Approaches', level=2)
    approaches = doc.add_paragraph('Traditional approaches to Blackjack include:')
    approaches.add_run('\n• Basic strategy tables')
    approaches.add_run('\n• Card counting systems')
    approaches.add_run('\n• Monte Carlo methods')
    approaches.add_run('\n• Rule-based expert systems')
    
    # How This Project Differs
    doc.add_heading('How This Project Differs', level=2)
    differences = doc.add_paragraph('This project differs from existing approaches by:')
    differences.add_run('\n• Using Q-learning for strategy development')
    differences.add_run('\n• Learning through experience rather than pre-programmed rules')
    differences.add_run('\n• Adapting to different game conditions')
    differences.add_run('\n• Providing real-time visualization of learning progress')
    
    # Add Methodology section
    doc.add_heading('4. Methodology', level=1)
    
    # Baseline Model
    doc.add_heading('Baseline Model', level=2)
    doc.add_paragraph('The Q-learning implementation uses a tabular approach with the following key components:')
    
    code = doc.add_paragraph()
    code.add_run('''class QLearningAgent:
    def __init__(self, learning_rate: float = 0.05, 
                 discount_factor: float = 0.95,
                 exploration_rate: float = 1.0,
                 exploration_decay: float = 0.9995):
        self.q_table = defaultdict(lambda: np.zeros(3))  # 3 actions: hit, stand, double
        self.lr = learning_rate
        self.gamma = discount_factor
        self.epsilon = exploration_rate
        self.epsilon_decay = exploration_decay
        self.min_epsilon = 0.005''')
    
    features = doc.add_paragraph('Key features:')
    features.add_run('\n• Adaptive exploration rate with decay')
    features.add_run('\n• Tabular Q-table for state-action values')
    features.add_run('\n• Three possible actions: hit, stand, double down')
    features.add_run('\n• Learning rate of 0.05 for stable updates')
    features.add_run('\n• Discount factor of 0.95 for future reward consideration')
    
    # Add Technical Implementation section
    doc.add_heading('5. Technical Implementation', level=1)
    
    # Environment Design
    doc.add_heading('Environment Design', level=2)
    doc.add_paragraph('The Blackjack environment implements casino rules:')
    
    code = doc.add_paragraph()
    code.add_run('''class BlackjackEnv:
    def __init__(self, num_decks: int = 4):
        self.deck = Deck(num_decks)
        self.player_hand: List[Card] = []
        self.dealer_hand: List[Card] = []
        self.game_over = False
        self.player_sum = 0
        self.dealer_sum = 0
        self.player_has_usable_ace = False
        self.dealer_has_usable_ace = False''')
    
    features = doc.add_paragraph('Features:')
    features.add_run('\n• Multiple deck support')
    features.add_run('\n• Proper ace handling')
    features.add_run('\n• Dealer strategy implementation')
    features.add_run('\n• Natural blackjack detection')
    
    # Add Results section
    doc.add_heading('6. Results', level=1)
    
    # Performance Metrics
    doc.add_heading('Performance Metrics', level=2)
    metrics = doc.add_paragraph('• Training convergence analysis')
    metrics.add_run('\n• Win rate progression')
    metrics.add_run('\n• Policy stability measures')
    metrics.add_run('\n• Computational efficiency')
    
    # Add Conclusions section
    doc.add_heading('7. Conclusions & Future Work', level=1)
    
    # Key Achievements
    doc.add_heading('Key Achievements', level=2)
    achievements = doc.add_paragraph('• Successful implementation of Q-learning for Blackjack')
    achievements.add_run('\n• Demonstrated learning capability')
    achievements.add_run('\n• Efficient state representation')
    achievements.add_run('\n• Practical training framework')
    
    # Add References section
    doc.add_heading('9. References', level=1)
    refs = doc.add_paragraph('1. Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction. MIT Press.')
    refs.add_run('\n2. Thorp, E. O. (1966). Beat the Dealer: A Winning Strategy for the Game of Twenty-One. Vintage.')
    refs.add_run('\n3. Mnih, V., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518(7540), 529-533.')
    refs.add_run('\n4. OpenAI Gym: A toolkit for developing and comparing reinforcement learning algorithms.')
    
    # Save the document
    doc.save('Blackjack_AI_Report.docx')

if __name__ == "__main__":
    create_word_report() 